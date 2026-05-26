import os
from docx import Document
from docx.shared import Pt

EXTENSIONS = (".reg", ".cpp", ".h", ".def")

def main():
    current_file = os.path.abspath(__file__)
    root_dir = os.path.dirname(current_file)

    doc = Document()
    doc.add_heading("Collected Source Files", level=0)

    for dirpath, dirnames, filenames in os.walk(root_dir):
        relative_dir = os.path.relpath(dirpath, root_dir)

        # Add a folder heading (skip root ".")
        if relative_dir != ".":
            doc.add_heading(f"Folder: {relative_dir}", level=1)

        for filename in sorted(filenames):
            if filename.endswith(EXTENSIONS):
                filepath = os.path.abspath(os.path.join(dirpath, filename))

                # Skip this script itself
                if filepath == current_file:
                    continue

                with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
                    code = f.read()

                doc.add_heading(filename, level=2)
                code_paragraph = doc.add_paragraph()
                code_run = code_paragraph.add_run(code)
                code_run.font.name = "Courier New"
                code_run.font.size = Pt(10)

    output_path = os.path.join(root_dir, "codes.docx")
    doc.save(output_path)
    print(f"Saved: {output_path}")

if __name__ == "__main__":
    main()
